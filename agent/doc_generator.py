"""
Document Generator – Phase 4

Produces a polished Microsoft Word (.docx) document using python-docx.
Renders document type, title, assumptions, task results, and reflection
in a professional business format.
"""

import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, Any

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger("agent.doc_generator")


class DocumentGenerator:
    """Generate a polished Word document from agent execution data."""

    PRIMARY_COLOR = RGBColor(0x1F, 0x45, 0x7E)   # Dark blue
    ACCENT_COLOR  = RGBColor(0x2E, 0x86, 0xAB)   # Mid blue
    LIGHT_GRAY    = RGBColor(0xF2, 0xF2, 0xF2)
    TEXT_COLOR    = RGBColor(0x33, 0x33, 0x33)

    def generate(
        self,
        document_type: str,
        title: str,
        request: str,
        plan: Dict[str, Any],
        execution_data: Dict[str, Any],
        reflection: str,
        filepath: str,
    ):
        doc = Document()
        self._configure_page(doc)
        self._apply_styles(doc)

        # ── Cover section ──────────────────────────────────────────────────────
        self._add_cover(doc, document_type, title)

        # ── Document metadata table ────────────────────────────────────────────
        self._add_metadata_table(doc, request, plan)

        # ── Assumptions (if any) ───────────────────────────────────────────────
        assumptions = plan.get("assumptions", [])
        if assumptions:
            self._add_heading(doc, "Assumptions & Clarifications", level=1)
            p = doc.add_paragraph()
            p.add_run(
                "The following assumptions were made for ambiguous or missing information in the request:"
            ).italic = True
            for assumption in assumptions:
                bullet = doc.add_paragraph(style="List Bullet")
                bullet.add_run(assumption)

        # ── Agent task list ────────────────────────────────────────────────────
        self._add_heading(doc, "Agent Task Plan", level=1)
        self._add_task_table(doc, plan["tasks"], execution_data)

        # ── Main content sections ─────────────────────────────────────────────
        self._add_heading(doc, "Document Content", level=1)
        sections = execution_data.get("sections", {})
        tasks = plan.get("tasks", [])

        for task in tasks:
            key = task.get("output_key", "")
            content = sections.get(key, "")
            if content:
                self._add_heading(doc, task["title"], level=2)
                self._add_content_paragraphs(doc, content)

        # ── Reflection / Quality Check ─────────────────────────────────────────
        self._add_heading(doc, "Agent Quality Review", level=1)
        p = doc.add_paragraph()
        p.add_run("Self-Check Assessment").bold = True
        doc.add_paragraph(reflection)

        # ── Footer note ───────────────────────────────────────────────────────
        self._add_footer(doc)

        doc.save(filepath)
        logger.info(f"Document saved: {filepath}")

    # ── Private helpers ────────────────────────────────────────────────────────

    def _configure_page(self, doc: Document):
        section = doc.sections[0]
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = section.right_margin = Inches(1)
        section.top_margin  = section.bottom_margin = Inches(1)

    def _apply_styles(self, doc: Document):
        styles = doc.styles

        # Normal style
        normal = styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        normal.font.color.rgb = self.TEXT_COLOR

        # Heading 1
        h1 = styles["Heading 1"]
        h1.font.name = "Calibri"
        h1.font.size = Pt(14)
        h1.font.bold = True
        h1.font.color.rgb = self.PRIMARY_COLOR
        h1.paragraph_format.space_before = Pt(18)
        h1.paragraph_format.space_after  = Pt(6)

        # Heading 2
        h2 = styles["Heading 2"]
        h2.font.name = "Calibri"
        h2.font.size = Pt(12)
        h2.font.bold = True
        h2.font.color.rgb = self.ACCENT_COLOR
        h2.paragraph_format.space_before = Pt(12)
        h2.paragraph_format.space_after  = Pt(4)

    def _add_cover(self, doc: Document, document_type: str, title: str):
        # Document type label
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(document_type.upper())
        run.font.size  = Pt(11)
        run.font.color.rgb = self.ACCENT_COLOR
        run.font.bold  = True
        run.font.name  = "Calibri"

        # Main title
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(title)
        run2.font.size  = Pt(22)
        run2.font.bold  = True
        run2.font.color.rgb = self.PRIMARY_COLOR
        run2.font.name  = "Calibri"

        # Date line
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(
            f"Generated by Autonomous AI Agent  •  {datetime.utcnow().strftime('%B %d, %Y')}"
        )
        run3.font.size  = Pt(10)
        run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run3.font.name  = "Calibri"
        run3.italic     = True

        doc.add_paragraph()  # spacer

        # Horizontal rule via paragraph border
        hr = doc.add_paragraph()
        self._add_bottom_border(hr, color="1F457E", size=12)
        doc.add_paragraph()  # spacer after rule

    def _add_metadata_table(self, doc: Document, request: str, plan: Dict):
        self._add_heading(doc, "Request Overview", level=1)

        table = doc.add_table(rows=4, cols=2)
        table.style = "Table Grid"
        table.columns[0].width = Inches(2)
        table.columns[1].width = Inches(5.5)

        rows_data = [
            ("Original Request",  request),
            ("Document Type",     plan["document_type"]),
            ("Document Title",    plan["document_title"]),
            ("Tasks Planned",     str(len(plan.get("tasks", [])))),
        ]

        for i, (label, value) in enumerate(rows_data):
            row = table.rows[i]
            # Label cell (shaded)
            label_cell = row.cells[0]
            label_cell.text = label
            label_cell.paragraphs[0].runs[0].bold = True
            label_cell.paragraphs[0].runs[0].font.color.rgb = self.PRIMARY_COLOR
            self._shade_cell(label_cell, "EBF0F8")

            # Value cell
            value_cell = row.cells[1]
            value_cell.text = value

        doc.add_paragraph()  # spacer

    def _add_task_table(self, doc: Document, tasks: list, execution_data: Dict):
        sections = execution_data.get("sections", {})
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        # Header row
        hdr_cells = table.rows[0].cells
        headers = ["#", "Task", "Description", "Status"]
        col_widths = [Inches(0.4), Inches(1.8), Inches(4.0), Inches(0.8)]

        for i, (cell, header, width) in enumerate(zip(hdr_cells, headers, col_widths)):
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            self._shade_cell(cell, "1F457E")
            cell.width = width

        # Data rows
        for task in tasks:
            key = task.get("output_key", "")
            status = "✓ Done" if key in sections else "✗ Failed"
            row = table.add_row()
            row.cells[0].text = str(task["id"])
            row.cells[1].text = task["title"]
            row.cells[2].text = task["description"]
            row.cells[3].text = status

            if "Done" in status:
                self._shade_cell(row.cells[3], "E8F5E9")
            else:
                self._shade_cell(row.cells[3], "FFEBEE")

        doc.add_paragraph()

    def _add_heading(self, doc: Document, text: str, level: int):
        style = f"Heading {level}"
        doc.add_heading(text, level=level)

    def _add_content_paragraphs(self, doc: Document, content: str):
        """Split content on blank lines and render each block."""
        blocks = content.split("\n\n")
        for block in blocks:
            block = block.strip()
            if not block:
                continue
            lines = block.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                # Detect markdown-style bullets
                if line.startswith("- ") or line.startswith("* "):
                    p = doc.add_paragraph(line[2:], style="List Bullet")
                elif line.startswith("**") and line.endswith("**"):
                    p = doc.add_paragraph()
                    p.add_run(line.strip("*")).bold = True
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                else:
                    doc.add_paragraph(line)

    def _add_footer(self, doc: Document):
        doc.add_paragraph()
        hr = doc.add_paragraph()
        self._add_bottom_border(hr, color="CCCCCC", size=6)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            "Generated by Autonomous AI Agent  •  Confidential  •  "
            + datetime.utcnow().strftime("%Y")
        )
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.italic = True

    def _shade_cell(self, cell, hex_color: str):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tc_pr.append(shd)

    def _add_bottom_border(self, paragraph, color="000000", size=6):
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(size))
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), color)
        pBdr.append(bottom)
        pPr.append(pBdr)
